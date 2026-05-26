from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path
import os
from datetime import datetime

class PDFConverter:
    """Handle PDF to Word conversions"""
    
    def __init__(self):
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
    
    def convert_simple(self, pdf_path: str) -> Path:
        """
        Simple conversion: Extract text using OCR and create Word document.
        """
        try:
            images = convert_from_path(pdf_path)
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('PDF Conversion', 0)
            title.alignment = 1  # Center alignment
            
            # Add metadata
            doc.add_paragraph(f"Original file: {Path(pdf_path).name}")
            doc.add_paragraph(f"Conversion date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Extract text from each page
            for i, image in enumerate(images):
                doc.add_heading(f"Page {i + 1}", level=1)
                
                # OCR to extract text
                text = pytesseract.image_to_string(image, lang='spa+eng')
                doc.add_paragraph(text)
                doc.add_page_break()
            
            # Save document
            output_path = self._get_output_path(pdf_path, "_simple.docx")
            doc.save(str(output_path))
            
            return output_path
        
        except Exception as e:
            raise Exception(f"Error in simple conversion: {str(e)}")
    
    def convert_advanced(self, pdf_path: str) -> Path:
        """
        Advanced conversion: Attempt to preserve formatting, tables, etc.
        """
        try:
            images = convert_from_path(pdf_path)
            
            doc = Document()
            
            # Add formatted title
            title = doc.add_heading('PDF Conversion - Advanced Mode', 0)
            title_format = title.runs[0]
            title_format.font.size = Pt(24)
            title_format.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add metadata with formatting
            info_para = doc.add_paragraph()
            info_para.add_run("Original file: ").bold = True
            info_para.add_run(Path(pdf_path).name)
            
            date_para = doc.add_paragraph()
            date_para.add_run("Conversion date: ").bold = True
            date_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            doc.add_paragraph()
            
            # Process each page
            for i, image in enumerate(images):
                page_heading = doc.add_heading(f"Page {i + 1}", level=1)
                page_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                
                # Extract text with OCR
                text = pytesseract.image_to_string(image, lang='spa+eng')
                
                # Try to detect and preserve structure
                lines = text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Simple heuristics for formatting
                        if len(line) > 50 and line.isupper():
                            para = doc.add_paragraph(line, style='Heading 2')
                        else:
                            para = doc.add_paragraph(line)
                        
                        # Improve readability
                        para_format = para.paragraph_format
                        para_format.space_after = Pt(6)
                        para_format.line_spacing = 1.15
                
                # Add image thumbnail of original page
                doc.add_paragraph()
                try:
                    # Resize image for thumbnail
                    thumb = image.copy()
                    thumb.thumbnail((400, 400))
                    
                    # Save temp image
                    temp_img_path = self.output_dir / f"temp_page_{i}.png"
                    thumb.save(str(temp_img_path))
                    
                    doc.add_picture(str(temp_img_path), width=Inches(4))
                    temp_img_path.unlink()
                except Exception as e:
                    doc.add_paragraph(f"[Could not embed page image: {str(e)}]")
                
                doc.add_page_break()
            
            # Save document
            output_path = self._get_output_path(pdf_path, "_advanced.docx")
            doc.save(str(output_path))
            
            return output_path
        
        except Exception as e:
            raise Exception(f"Error in advanced conversion: {str(e)}")
    
    def _get_output_path(self, pdf_path: str, suffix: str) -> Path:
        """Generate output file path"""
        pdf_name = Path(pdf_path).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{pdf_name}_{timestamp}{suffix}"
        return self.output_dir / filename
